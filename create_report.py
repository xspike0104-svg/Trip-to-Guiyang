from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Set default font for Chinese
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Title
title = doc.add_heading('PT HEBANG BIOTECHNOLOGY INDONESIA', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('印尼法律资讯第十一期影响分析报告', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'报告日期：2026年4月16日')
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Company Overview
doc.add_heading('一、公司概况', level=2)
table1 = doc.add_table(rows=5, cols=2)
table1.style = 'Table Grid'
data1 = [
    ('公司名称', 'PT HEBANG BIOTECHNOLOGY INDONESIA'),
    ('母公司', '四川和邦生物科技股份有限公司（A股上市公司，603077.SH）'),
    ('项目位置', '印尼东爪哇 JIIPE 经济特区（Java International Industrial Port and Estate）'),
    ('产品', '草甘膦（Glyphosate）'),
    ('规模', '35万吨/年（已从20万吨/年调整，为全球最大草甘膦单体装置之一）'),
]
for i, (key, val) in enumerate(data1):
    table1.rows[i].cells[0].text = key
    table1.rows[i].cells[1].text = val

doc.add_paragraph()

# Section 1 - TKDN
doc.add_heading('二、TKDN 本地成分含量规定', level=2)
doc.add_heading('2.1 规定概述', level=3)
p = doc.add_paragraph()
p.add_run('TKDN（Tingkat Komponen Dalam Negeri）').bold = True
p.add_run(' 是印尼最重要的工业政策工具，旨在推动本地化产业发展。根据2026年最新政策，化工行业的本地成分要求平均达到 ')
p.add_run('45%').bold = True
p.add_run('。')

doc.add_heading('2.2 对项目的影响', level=3)
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
data2 = [
    ('风险等级', '⚠️ 高风险'),
    ('TKDN 计算方式', '原材料（25%）+ 加工成本（15%）+ 直接人工（10%）+ 其他'),
    ('最低要求', '化工行业平均 45%'),
    ('违规后果', '无法享受税收优惠、可能被取消投资激励资格'),
    ('建议行动', '提前与印尼工业部确认 TKDN 计算方法，准备本地化方案'),
]
for i, (key, val) in enumerate(data2):
    table2.rows[i].cells[0].text = key
    table2.rows[i].cells[1].text = val

doc.add_paragraph()

# Section 2 - KIH
doc.add_heading('三、重点工业企业（KIH）专项许可证', level=2)
doc.add_heading('3.1 KIH 资格说明', level=3)
p = doc.add_paragraph()
p.add_run('KIH（Kelompok Industri Harapan）').bold = True
p.add_run(' 是印尼政府为促进重点工业发展而设立的专项许可制度。化工行业已被列为2026年优先发展的40个行业之一。')

doc.add_heading('3.2 你的项目适用性', level=3)
p = doc.add_paragraph()
p.add_run('35万吨/年草甘膦项目').bold = True
p.add_run(' 完全符合"大规模化工项目"的定义，可申请 KIH 资格认定。')

doc.add_heading('3.3 可享受的优惠', level=3)
benefits = [
    '简化许可证审批流程',
    '企业所得税减免优惠',
    '优先获得工业用地',
    '进口设备关税减免',
    '简化环境影响评估（EIA）流程',
]
for b in benefits:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('3.4 建议行动', level=3)
p = doc.add_paragraph()
p.add_run('尽快向印尼工业部（Ministry of Industry）提交 KIH 资格申请').bold = True
p.add_run('，享受审批便利和税收减免。')

doc.add_paragraph()

# Section 3 - JIIPE
doc.add_heading('四、JIIPE 工业园区政策', level=2)
doc.add_heading('4.1 园区优势', level=3)
p = doc.add_paragraph('你的项目位于 JIIPE 经济特区，可享受以下优势：')

jiipe_benefits = [
    '税收优惠延长至2026年',
    '企业所得税减免',
    '进口关税豁免或减免',
    '简化海关流程',
    '完善的基础设施配套',
    '一站式投资服务',
]
for b in jiipe_benefits:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('4.2 税收激励详情', level=3)
table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Table Grid'
data3 = [
    ('激励类型', '说明'),
    ('企业所得税优惠', '符合条件的企业可享受额外减免'),
    ('进口关税优惠', '用于生产的设备、原材料进口关税减免'),
    ('增值税优惠', '本地采购可享受增值税优惠'),
]
for i, (key, val) in enumerate(data3):
    table3.rows[i].cells[0].text = key
    table3.rows[i].cells[1].text = val

doc.add_paragraph()

# Section 4 - Clean Energy
doc.add_heading('五、清洁能源与环保要求', level=2)
doc.add_heading('5.1 清洁能源趋势', level=3)
p = doc.add_paragraph('印尼政府正在推动清洁能源转型，大型工业项目需符合环保标准才能获得批准。')

doc.add_heading('5.2 对项目的具体要求', level=3)
env_requirements = [
    '环境影响评估（EIA）：必须通过并获得批准',
    '废水排放标准：需符合印尼国家废水排放标准',
    '废气控制：粉尘和挥发性有机物排放需达标',
    '清洁能源使用：鼓励使用可再生能源',
    '定期环境监测：需每6个月提交环境监测报告',
]
for r in env_requirements:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('5.3 风险提示', level=3)
p = doc.add_paragraph()
p.add_run('⚠️ 注意：').bold = True
p.add_run('印尼最高法院2024年已多次判决撤销环境许可证，淡水河谷镍矿项目因林区许可问题已于2026年1月暂停。务必确保 EIA 审批合规。')

doc.add_paragraph()

# Section 5 - Other Regulations
doc.add_heading('六、其他重要法规提示', level=2)

doc.add_heading('6.1 限定继承税及股息税', level=3)
p = doc.add_paragraph('根据法律资讯第二部分，印尼正在推进限定继承税改革，可能影响利润分配时的税务处理。建议关注股息汇出时的预扣税规定。')

doc.add_heading('6.2 平台经济法规', level=3)
p = doc.add_paragraph('如项目涉及数字化采购或销售平台，需遵守印尼数字经济相关法规。')

doc.add_heading('6.3 劳工规定', level=3)
p = doc.add_paragraph('KIH 企业需遵守印尼劳工法规，包括最低工资、就业许可、工作环境安全标准等。')

doc.add_paragraph()

# Summary Table
doc.add_heading('七、行动建议汇总', level=2)
table4 = doc.add_table(rows=6, cols=3)
table4.style = 'Table Grid'
headers4 = ['优先级', '行动项', '理由']
for i, h in enumerate(headers4):
    table4.rows[0].cells[i].text = h
    table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data4 = [
    ('🔴 高', '申请 KIH 资格认定', '享受许可证便利和税收减免'),
    ('🔴 高', '确认 TKDN 计算方案', '避免不符合导致的项目延误或罚款'),
    ('🟡 中', '利用 JIIPE 园区优惠', '确保享受所有税收减免政策'),
    ('🟡 中', '规划清洁能源方案', '符合印尼环保趋势，避免 EIA 审批问题'),
    ('🟢 低', '关注股息税政策', '利润分配时注意预扣税影响'),
]
for i, (p_level, action, reason) in enumerate(data4, 1):
    table4.rows[i].cells[0].text = p_level
    table4.rows[i].cells[1].text = action
    table4.rows[i].cells[2].text = reason

doc.add_paragraph()

# Disclaimer
doc.add_heading('八、免责声明', level=2)
p = doc.add_paragraph()
p.add_run('本报告基于公开信息和印尼法律资讯第十一期内容编制，仅供参考。具体法规要求请咨询印尼当地专业律师或合规顾问。').font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('资料来源：大成Dentons律师事务所《印尼法律资讯第十一期》（2026年3月20日）').font.size = Pt(9)

# Save
output_path = r'C:\Users\pc\Desktop\印尼法律分析_PT_Hebang_Biotechnology.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
